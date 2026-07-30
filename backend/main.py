# import os
# import io
# import re
# from fastapi import FastAPI, UploadFile, File, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import StreamingResponse
# from pydantic import BaseModel
# from google import genai
# from docx import Document
# from fpdf import FPDF
# from PIL import Image
# from dotenv import load_dotenv
# from pdf2image import convert_from_bytes
# from fastapi import FastAPI, UploadFile, File, Form, HTTPException
# from typing import Optional
# from fastapi import FastAPI, UploadFile, File, Form, HTTPException
# from pydantic import BaseModel

# load_dotenv()

# app = FastAPI(title="Handwriting Transcriber API (Gemini)")

# # Enable CORS and expose headers so the browser can download files properly
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
#     expose_headers=["Content-Disposition"],
# )

# # Initialize Google Gemini Client
# client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

# class ExportRequest(BaseModel):
#     text: str
#     format: str  # "docx" or "pdf"

# def clean_transcription_text(text: str) -> str:
#     """Strips out model thinking blocks (<think>...</think>) and conversational preambles."""
#     if not text:
#         return ""
    
#     # 1. Remove <think>...</think> blocks including unclosed <think> tags
#     cleaned = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL | re.IGNORECASE)
#     cleaned = re.sub(r'<think>.*$', '', cleaned, flags=re.DOTALL | re.IGNORECASE)
    
#     # 2. Split lines and filter out introductory fluff
#     lines = cleaned.strip().split("\n")
#     filtered_lines = []
#     skip_intro = True
    
#     for line in lines:
#         line_str = line.strip()
#         if skip_intro:
#             lower = line_str.lower()
#             if (
#                 lower.startswith("here is") or 
#                 lower.startswith("here's") or 
#                 lower.startswith("the handwritten text") or 
#                 lower.startswith("transcription:") or 
#                 not line_str
#             ):
#                 continue
#             else:
#                 skip_intro = False
#         filtered_lines.append(line)
        
#     result = "\n".join(filtered_lines).strip()
#     return result if result else cleaned.strip()

# @app.post("/api/transcribe")
# async def transcribe_file(
#     file: UploadFile = File(...),
#     language: Optional[str] = Form("English")  # Default to English
# ):
#     try:
#         file_bytes = await file.read()
        
#         # Pass the language context into your OCR/Vision prompt
#         # Example prompt structure for Gemini/GPT-4 Vision:
#         prompt = f"""
#         Transcribe the handwritten text in this image accurately. 
#         The expected primary language of the text is {language}.
#         Preserve original formatting, paragraphs, and line breaks where possible.
#         Return ONLY the transcribed text.
#         """
        
#         # Call your AI Vision function here using `prompt` and `file_bytes`
#         transcribed_text = await run_ocr_vision_model(file_bytes, file.filename, prompt)

#         return {"text": transcribed_text, "language": language}

#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# @app.post("/api/export")
# async def export_document(data: ExportRequest):
#     # Sanitize text before exporting
#     cleaned_text = clean_transcription_text(data.text)

#     if data.format == "docx":
#         doc = Document()
#         doc.add_heading("Transcribed Document", level=1)
#         for line in cleaned_text.split("\n"):
#             if line.strip():
#                 doc.add_paragraph(line)
        
#         file_stream = io.BytesIO()
#         doc.save(file_stream)
#         file_stream.seek(0)
        
#         return StreamingResponse(
#             file_stream,
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             headers={"Content-Disposition": "attachment; filename=transcription.docx"}
#         )

#     elif data.format == "pdf":
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_font("Helvetica", size=12)
#         for line in cleaned_text.split("\n"):
#             pdf.multi_cell(0, 8, txt=line.encode('latin-1', 'replace').decode('latin-1'))
#             pdf.ln(1)
            
#         pdf_bytes = pdf.output()
#         file_stream = io.BytesIO(pdf_bytes)
        
#         return StreamingResponse(
#             file_stream,
#             media_type="application/pdf",
#             headers={"Content-Disposition": "attachment; filename=transcription.pdf"}
#         )

#     raise HTTPException(status_code=400, detail="Invalid format selected")


# class CleanupRequest(BaseModel):
#     text: str

# @app.post("/api/cleanup")
# async def cleanup_text(payload: CleanupRequest):
#     if not payload.text.strip():
#         raise HTTPException(status_code=400, detail="Text cannot be empty.")

#     try:
#         prompt = f"""
#         You are an expert editor and proofreader. Clean up the following OCR-extracted text from handwriting.
#         - Fix spelling errors and misread words caused by messy handwriting.
#         - Fix grammatical issues, punctuation, and sentence structures.
#         - Preserve the original meaning and tone completely.
#         - Do not add conversational commentary. Return ONLY the cleaned-up text.

#         Text to clean:
#         {payload.text}
#         """

#         # Call your LLM model here (e.g., Gemini 1.5/2.0, GPT-4o, etc.)
#         # Example pseudo-call:
#         # response = model.generate_content(prompt)
#         # cleaned_text = response.text.strip()
        
#         cleaned_text = await run_llm_cleanup(prompt) # Replace with your model caller

#         return {"cleaned_text": cleaned_text}

#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))



import asyncio
import io
import json
import os
from typing import List, Optional

from dotenv import load_dotenv
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fpdf import FPDF
from google import genai
from google.genai import types
from pdf2image import convert_from_bytes
from pydantic import BaseModel

# Load environment variables from .env file
load_dotenv()

app = FastAPI(title="InkSync API")

# Enable CORS for frontend connection
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Gemini Client (Ensure GEMINI_API_KEY is set in your environment)
gemini_client = genai.Client()


class CleanupRequest(BaseModel):
    text: str


class ExportRequest(BaseModel):
    text: str
    format: str


class AnalyzeRequest(BaseModel):
    text: str
    prompt: Optional[str] = None


def _clean_transcription_text(text: str) -> str:
    if not text:
        return ""
    cleaned = text.strip()
    return cleaned.replace("Here is the cleaned text:", "").strip()


def _call_gemini_text(prompt: str) -> str:
    response = gemini_client.models.generate_content(model="gemini-3.6-flash", contents=prompt)
    return response.text.strip() if response.text else ""


def _call_gemini_json(prompt: str) -> dict:
    response = gemini_client.models.generate_content(model="gemini-3.6-flash", contents=prompt)
    raw_text = response.text.strip() if response.text else "{}"
    try:
        return json.loads(raw_text)
    except json.JSONDecodeError:
        return {"category": "General", "tags": ["General"]}


async def _transcribe_single_bytes(file_bytes: bytes, mime_type: str, prompt: str) -> str:
    def _run() -> str:
        image_part = types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
        response = gemini_client.models.generate_content(model="gemini-3.6-flash", contents=[image_part, prompt])
        return response.text.strip() if response.text else ""

    return await asyncio.to_thread(_run)


async def _extract_tags_and_category(text: str) -> dict:
    prompt = f"""
    Analyze this handwritten note transcription and return ONLY valid JSON in this shape:
    {{"category": "...", "tags": ["...", "..."]}}
    Choose a short category and 3-5 relevant tags.

    Text:
    {text}
    """
    return await asyncio.to_thread(_call_gemini_json, prompt)


@app.post("/api/transcribe")
async def transcribe_file(
    files: List[UploadFile] = File(...),
    language: Optional[str] = Form("English")
):
    if not files:
        raise HTTPException(status_code=400, detail="At least one file is required.")

    try:
        uploads: List[tuple[str, bytes, str]] = []
        for uploaded_file in files:
            file_bytes = await uploaded_file.read()
            mime_type = uploaded_file.content_type or "image/jpeg"

            if mime_type == "application/pdf":
                try:
                    pages = convert_from_bytes(file_bytes)
                except Exception:
                    pages = []

                if not pages:
                    raise HTTPException(status_code=400, detail="Unable to read PDF pages. Please upload a valid PDF.")

                for index, page in enumerate(pages):
                    buffer = io.BytesIO()
                    page.save(buffer, format="JPEG")
                    uploads.append((f"{uploaded_file.filename} page {index + 1}", buffer.getvalue(), "image/jpeg"))
            else:
                uploads.append((uploaded_file.filename, file_bytes, mime_type))

        prompt_template = f"""
        Transcribe the handwritten text in this document as accurately as possible.
        The expected primary language of the text is {language}.
        Preserve the layout, line breaks, and paragraph structures where sensible.
        Do NOT add any conversational preambles, intros, or explanations.
        Return ONLY the raw extracted transcribed text.
        """

        tasks = [
            asyncio.create_task(_transcribe_single_bytes(file_bytes, mime_type, prompt_template))
            for _, file_bytes, mime_type in uploads
        ]
        transcribed_parts = await asyncio.gather(*tasks)

        combined_text = "\n\n".join(
            transcribed
            for transcribed in transcribed_parts
            if transcribed.strip()
        )

        if not combined_text.strip():
            raise HTTPException(status_code=500, detail="No readable text was extracted.")

        analysis = await _extract_tags_and_category(combined_text)

        return {
            "text": combined_text,
            "language": language,
            "source_count": len(uploads),
            "tags": analysis.get("tags", ["General"]),
            "category": analysis.get("category", "General"),
        }

    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(exc)}")


@app.post("/api/cleanup")
async def cleanup_text(payload: CleanupRequest):
    if not payload.text.strip():
        raise HTTPException(status_code=400, detail="Text cannot be empty.")

    try:
        prompt = f"""
        You are an expert editor. Clean up and refine the following OCR-extracted handwritten text:
        - Correct spelling mistakes, mistranscribed characters, and handwriting OCR typos.
        - Fix grammatical errors, punctuation, and sentence structures.
        - Preserve the original meaning and tone completely.
        - Do NOT add any conversational preambles (e.g. "Here is the cleaned text:").
        - Return ONLY the polished, cleaned-up text.

        Text to clean:
        {payload.text}
        """

        cleaned_text = await asyncio.to_thread(_call_gemini_text, prompt)
        return {"cleaned_text": _clean_transcription_text(cleaned_text)}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cleanup failed: {str(e)}")


@app.post("/api/analyze")
async def analyze_notes(payload: AnalyzeRequest):
    if not payload.text.strip():
        raise HTTPException(status_code=400, detail="Text cannot be empty.")

    try:
        prompt = f"""
        Answer the user's question using the supplied notes.
        Keep the answer concise, practical, and useful.

        Notes:
        {payload.text}

        User question:
        {payload.prompt or 'Summarize these notes into 3 key action items.'}
        """
        answer = await asyncio.to_thread(_call_gemini_text, prompt)
        return {"answer": answer}

    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(exc)}")


@app.post("/api/export")
async def export_file(payload: ExportRequest):
    if not payload.text.strip():
        raise HTTPException(status_code=400, detail="Text cannot be empty.")

    format_type = payload.format.lower()

    if format_type == "docx":
        doc = Document()
        doc.add_heading("InkSync Transcription", level=1)
        doc.add_paragraph(payload.text)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=transcription.docx"},
        )

    elif format_type == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 8, payload.text)

        buffer = io.BytesIO()
        pdf_output = pdf.output()
        buffer.write(pdf_output)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=transcription.pdf"},
        )

    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use 'docx' or 'pdf'.")